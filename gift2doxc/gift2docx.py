import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import re
import random
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os
import collections


class GiftToDocxConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("Convertidor de GIFT a DOCX")
        self.root.geometry("800x600")
        self.root.resizable(True, True)

        # Variables
        self.gift_files = []  # Lista de archivos seleccionados
        self.output_dir = tk.StringVar()
        self.current_step = 0
        self.questions = []
        self.answer_keys = {}
        self.selected_question_index = tk.IntVar(value=0)

        # Crear frames para cada paso
        self.frames = []
        for _ in range(5):
            frame = ttk.Frame(self.root, padding="10")
            frame.grid(row=0, column=0, sticky="nsew")
            self.frames.append(frame)

        # Paso 1: Seleccionar archivos GIFT
        self.setup_step1()

        # Paso 2: Opciones de configuración
        self.setup_step2()

        # Paso 3: Vista previa interactiva
        self.setup_step3()

        # Paso 4: configuración pagina
        self.setup_step4()

        # Paso 5: Resumen y finalización
        self.setup_step5()

        # Configurar los botones de navegación
        self.setup_navigation()

        # Mostrar el primer paso
        self.show_step(0)

        # Configurar grid
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)

    def setup_step1(self):
        frame = self.frames[0]

        # Título
        ttk.Label(frame, text="Paso 1: Seleccionar archivos GIFT", font=("Arial", 14, "bold")).grid(row=0, column=0, columnspan=3, pady=10, sticky="w")

        # Instrucciones
        ttk.Label(frame, text="Seleccione uno o varios archivos en formato GIFT:", wraplength=500).grid(row=1, column=0, columnspan=3, pady=10, sticky="w")

        # Lista de archivos seleccionados
        ttk.Label(frame, text="Archivos seleccionados:").grid(row=2, column=0, sticky="w", padx=5)

        self.files_listbox = tk.Listbox(frame, height=6)
        self.files_listbox.grid(row=3, column=0, columnspan=2, padx=5, pady=5, sticky="ew")

        # Scrollbar para la lista
        files_scrollbar = ttk.Scrollbar(frame, orient="vertical", command=self.files_listbox.yview)
        files_scrollbar.grid(row=3, column=2, sticky="ns", pady=5)
        self.files_listbox.configure(yscrollcommand=files_scrollbar.set)

        # Botones para manejo de archivos
        buttons_frame = ttk.Frame(frame)
        buttons_frame.grid(row=4, column=0, columnspan=3, pady=10, sticky="ew")

        ttk.Button(buttons_frame, text="Agregar archivos...", command=self.add_gift_files).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="Quitar seleccionado", command=self.remove_selected_file).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="Limpiar todo", command=self.clear_all_files).pack(side=tk.LEFT, padx=5)

        # Configurar grid
        frame.columnconfigure(0, weight=1)

    def setup_step2(self):
        frame = self.frames[1]

        # Título
        ttk.Label(frame, text="Paso 2: Opciones de configuración", font=("Arial", 14, "bold")).grid(row=0, column=0, columnspan=2, pady=10, sticky="w")

        # Directorio de salida
        ttk.Label(frame, text="Directorio para guardar los archivos de salida:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        ttk.Entry(frame, textvariable=self.output_dir, width=50).grid(row=2, column=0, padx=5, pady=5, sticky="we")
        ttk.Button(frame, text="Examinar...", command=self.browse_output_dir).grid(row=2, column=1, padx=5, pady=5)

        # Opciones de formato
        ttk.Label(frame, text="Opciones de formato:").grid(row=3, column=0, columnspan=2, padx=5, pady=10, sticky="w")

        # Orden de preguntas
        ttk.Label(frame, text="Orden de las preguntas:").grid(row=4, column=0, padx=20, pady=5, sticky="w")
        self.question_order = tk.StringVar(value="por_archivos")
        ttk.Radiobutton(frame, text="Mantener orden por archivos", variable=self.question_order, value="por_archivos").grid(row=5, column=0, padx=40, pady=2, sticky="w")
        ttk.Radiobutton(frame, text="Aleatorio (mezclar todas)", variable=self.question_order, value="aleatorio").grid(row=6, column=0, padx=40, pady=2, sticky="w")

        # Opciones de respuestas
        self.randomize_options = tk.BooleanVar(value=True)
        ttk.Checkbutton(frame, text="Aleatorizar opciones de respuesta", variable=self.randomize_options).grid(row=7, column=0, padx=20, pady=5, sticky="w")

        # Umbral de detección de problemas
        ttk.Label(frame, text="Detección de problemas:").grid(row=8, column=0, padx=20, pady=(10,5), sticky="w")
        umbral_frame = ttk.Frame(frame)
        umbral_frame.grid(row=9, column=0, padx=40, pady=2, sticky="w")

        ttk.Label(umbral_frame, text="Umbral diferencia de caracteres:").pack(side=tk.LEFT)
        self.threshold_var = tk.IntVar(value=10)
        threshold_spinbox = tk.Spinbox(umbral_frame, from_=5, to=50, width=5, textvariable=self.threshold_var)
        threshold_spinbox.pack(side=tk.LEFT, padx=5)
        ttk.Label(umbral_frame, text="caracteres").pack(side=tk.LEFT)

        # Codificación de archivos
        ttk.Label(frame, text="Codificación de archivos:").grid(row=10, column=0, padx=20, pady=(10,5), sticky="w")
        self.encoding_var = tk.StringVar(value="auto")
        encodings_frame = ttk.Frame(frame)
        encodings_frame.grid(row=11, column=0, padx=40, pady=2, sticky="w")

        ttk.Radiobutton(encodings_frame, text="Detectar automáticamente", variable=self.encoding_var, value="auto").pack(anchor="w")
        ttk.Radiobutton(encodings_frame, text="UTF-8", variable=self.encoding_var, value="utf-8").pack(anchor="w")
        ttk.Radiobutton(encodings_frame, text="Latin-1 (ISO-8859-1)", variable=self.encoding_var, value="latin-1").pack(anchor="w")
        ttk.Radiobutton(encodings_frame, text="Windows-1252", variable=self.encoding_var, value="cp1252").pack(anchor="w")

        # Configurar grid
        frame.columnconfigure(0, weight=1)

    def setup_step3(self):
        frame = self.frames[2]

        # Título
        ttk.Label(frame, text="Paso 3: Vista previa y edición", font=("Arial", 14, "bold")).grid(row=0, column=0, columnspan=2, pady=10, sticky="w")

        # Frame principal con dos paneles
        main_frame = ttk.Frame(frame)
        main_frame.grid(row=1, column=0, columnspan=2, sticky="nsew", padx=5, pady=5)
        main_frame.columnconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=2)
        main_frame.rowconfigure(0, weight=1)

        # Panel izquierdo: Lista de preguntas y controles
        left_panel = ttk.LabelFrame(main_frame, text="Lista de preguntas", padding="5")
        left_panel.grid(row=0, column=0, sticky="nsew", padx=(0,5))
        left_panel.columnconfigure(0, weight=1)
        left_panel.rowconfigure(1, weight=1)  # La lista debe expandirse
        left_panel.rowconfigure(0, weight=0)  # Botones superiores fijos
        left_panel.rowconfigure(2, weight=0)  # Botones inferiores fijos

        # Botones de control superiores
        control_frame = ttk.Frame(left_panel)
        control_frame.grid(row=0, column=0, columnspan=2, sticky="ew", pady=(0,5))
        
        ttk.Button(control_frame, text="Cargar GIFT", command=self.process_gift_files).pack(side=tk.LEFT, padx=2)
        ttk.Button(control_frame, text="+ Alternativas", command=self.add_multiple_choice_question).pack(side=tk.LEFT, padx=2)
        ttk.Button(control_frame, text="+ Desarrollo", command=self.add_essay_question).pack(side=tk.LEFT, padx=2)

        # Listbox para preguntas con scrollbar
        self.questions_listbox = tk.Listbox(left_panel, height=12)
        self.questions_listbox.grid(row=1, column=0, sticky="nsew", padx=(0,5), pady=5)
        self.questions_listbox.bind("<<ListboxSelect>>", self.on_question_select)

        questions_scrollbar = ttk.Scrollbar(left_panel, orient="vertical", command=self.questions_listbox.yview)
        questions_scrollbar.grid(row=1, column=1, sticky="ns", pady=5)
        self.questions_listbox.configure(yscrollcommand=questions_scrollbar.set)

        # Botones de orden y eliminación (ahora más visibles)
        order_frame = ttk.Frame(left_panel)
        order_frame.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(5,0))
        
        ttk.Button(order_frame, text="↑ Subir", command=self.move_question_up).pack(side=tk.LEFT, padx=2)
        ttk.Button(order_frame, text="↓ Bajar", command=self.move_question_down).pack(side=tk.LEFT, padx=2)
        ttk.Button(order_frame, text="Eliminar", command=self.delete_question).pack(side=tk.LEFT, padx=2)

        # Panel derecho: Editor de pregunta
        right_panel = ttk.LabelFrame(main_frame, text="Editor de pregunta", padding="5")
        right_panel.grid(row=0, column=1, sticky="nsew")
        right_panel.columnconfigure(0, weight=1)
        right_panel.rowconfigure(1, weight=1)

        # Información de la pregunta actual
        self.question_info_label = ttk.Label(right_panel, text="Seleccione una pregunta para editar", font=("Arial", 10, "bold"))
        self.question_info_label.grid(row=0, column=0, sticky="w", pady=5)

        # Frame scrollable para el editor
        self.editor_canvas = tk.Canvas(right_panel, highlightthickness=0)
        self.editor_scrollbar = ttk.Scrollbar(right_panel, orient="vertical", command=self.editor_canvas.yview)
        self.editor_frame = ttk.Frame(self.editor_canvas)

        self.editor_canvas.grid(row=1, column=0, sticky="nsew")
        self.editor_scrollbar.grid(row=1, column=1, sticky="ns")
        self.editor_canvas.configure(yscrollcommand=self.editor_scrollbar.set)

        self.editor_canvas.create_window((0, 0), window=self.editor_frame, anchor="nw")
        self.editor_frame.bind("<Configure>", lambda e: self.editor_canvas.configure(scrollregion=self.editor_canvas.bbox("all")))

        # Botones de acción inferiores
        buttons_frame = ttk.Frame(frame)
        buttons_frame.grid(row=2, column=0, columnspan=2, pady=10, sticky="ew")

        # Frame para limpieza masiva
        cleanup_frame = ttk.LabelFrame(buttons_frame, text="Limpieza masiva", padding="5")
        cleanup_frame.pack(side=tk.LEFT, padx=5, fill="y")

        ttk.Label(cleanup_frame, text="Eliminar texto:").pack(side=tk.LEFT)
        self.cleanup_text = tk.StringVar()
        cleanup_entry = ttk.Entry(cleanup_frame, textvariable=self.cleanup_text, width=15)
        cleanup_entry.pack(side=tk.LEFT, padx=2)
        ttk.Button(cleanup_frame, text="Aplicar", command=self.apply_mass_cleanup).pack(side=tk.LEFT, padx=2)

        # Botones comunes predefinidos
        common_frame = ttk.Frame(cleanup_frame)
        common_frame.pack(side=tk.LEFT, padx=5)
        ttk.Button(common_frame, text="[html]", command=lambda: self.quick_cleanup("[html]")).pack(side=tk.TOP, pady=1)
        ttk.Button(common_frame, text="[moodle]", command=lambda: self.quick_cleanup("[moodle]")).pack(side=tk.TOP, pady=1)

        # Frame para puntajes
        score_frame = ttk.LabelFrame(buttons_frame, text="Puntajes", padding="5")
        score_frame.pack(side=tk.LEFT, padx=5, fill="y")

        ttk.Label(score_frame, text="Puntaje:").pack(side=tk.LEFT)
        self.mass_score = tk.StringVar(value="1")
        score_entry = ttk.Entry(score_frame, textvariable=self.mass_score, width=5)
        score_entry.pack(side=tk.LEFT, padx=2)

        score_buttons = ttk.Frame(score_frame)
        score_buttons.pack(side=tk.LEFT, padx=5)
        ttk.Button(score_buttons, text="A todas", command=self.apply_score_to_all).pack(side=tk.TOP, pady=1)
        ttk.Button(score_buttons, text="A alternativas", command=lambda: self.apply_score_by_type("multiple_choice")).pack(side=tk.TOP, pady=1)
        ttk.Button(score_buttons, text="A desarrollo", command=lambda: self.apply_score_by_type("essay")).pack(side=tk.TOP, pady=1)

        # Botones finales
        ttk.Button(buttons_frame, text="Guardar como GIFT", command=self.save_as_gift).pack(side=tk.RIGHT, padx=5)
        ttk.Button(buttons_frame, text="Procesar y continuar", command=self.apply_randomization).pack(side=tk.RIGHT, padx=5)

        # Configurar grid del frame principal
        frame.columnconfigure(0, weight=1)
        frame.rowconfigure(1, weight=1)

    def setup_step4(self):
        frame = self.frames[3]

        # Título
        ttk.Label(frame, text="Paso 4: Configuración de página", font=("Arial", 14, "bold")).grid(row=0, column=0, columnspan=2, pady=10, sticky="w")

        # Crear un canvas con scrollbar para manejar el contenido
        canvas = tk.Canvas(frame)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.grid(row=1, column=0, sticky="nsew", padx=5, pady=5)
        scrollbar.grid(row=1, column=1, sticky="ns", pady=5)

        # Variables de configuración
        self.exam_title = tk.StringVar(value="EXAMEN")
        self.answer_sheet_title = tk.StringVar(value="HOJA DE RESPUESTAS")
        self.institution_name = tk.StringVar(value="")
        self.course_name = tk.StringVar(value="")
        self.exam_date = tk.StringVar(value="")
        
        # Campos de estudiante
        self.show_student_name = tk.BooleanVar(value=True)
        self.show_student_id = tk.BooleanVar(value=True)
        self.show_student_section = tk.BooleanVar(value=False)
        self.show_exam_score = tk.BooleanVar(value=True)
        
        # Configuración de fuentes
        self.title_font_size = tk.IntVar(value=16)
        self.question_font_size = tk.IntVar(value=11)
        self.font_name = tk.StringVar(value="Arial")
        
        # Configuración de página
        self.page_size = tk.StringVar(value="A4")
        self.margin_top = tk.DoubleVar(value=2.5)
        self.margin_bottom = tk.DoubleVar(value=2.5)
        self.margin_left = tk.DoubleVar(value=2.5)
        self.margin_right = tk.DoubleVar(value=2.5)
        
        # Logo (esquema básico)
        self.show_logo = tk.BooleanVar(value=False)
        self.logo_position = tk.StringVar(value="top_left")
        
        # Configuración de respuestas correctas
        self.show_statistics = tk.BooleanVar(value=True)
        self.show_detailed_info = tk.BooleanVar(value=True)
        self.show_file_info = tk.BooleanVar(value=True)

        # Configurar grid principal para dos columnas
        scrollable_frame.columnconfigure(0, weight=1)
        scrollable_frame.columnconfigure(1, weight=1)

        # COLUMNA IZQUIERDA
        left_column = ttk.Frame(scrollable_frame)
        left_column.grid(row=0, column=0, sticky="nsew", padx=(0, 10))

        # COLUMNA DERECHA
        right_column = ttk.Frame(scrollable_frame)
        right_column.grid(row=0, column=1, sticky="nsew", padx=(10, 0))

        # === COLUMNA IZQUIERDA ===
        row_left = 0

        # === INFORMACIÓN GENERAL ===
        general_frame = ttk.LabelFrame(left_column, text="Información General", padding="10")
        general_frame.grid(row=row_left, column=0, sticky="ew", pady=5)
        general_frame.columnconfigure(1, weight=1)

        ttk.Label(general_frame, text="Título del examen:").grid(row=0, column=0, sticky="w", padx=5, pady=2)
        ttk.Entry(general_frame, textvariable=self.exam_title, width=30).grid(row=0, column=1, sticky="ew", padx=5, pady=2)

        ttk.Label(general_frame, text="Institución:").grid(row=1, column=0, sticky="w", padx=5, pady=2)
        ttk.Entry(general_frame, textvariable=self.institution_name, width=30).grid(row=1, column=1, sticky="ew", padx=5, pady=2)

        ttk.Label(general_frame, text="Curso/Materia:").grid(row=2, column=0, sticky="w", padx=5, pady=2)
        ttk.Entry(general_frame, textvariable=self.course_name, width=30).grid(row=2, column=1, sticky="ew", padx=5, pady=2)

        ttk.Label(general_frame, text="Fecha del examen:").grid(row=3, column=0, sticky="w", padx=5, pady=2)
        ttk.Entry(general_frame, textvariable=self.exam_date, width=30).grid(row=3, column=1, sticky="ew", padx=5, pady=2)

        row_left += 1

        # === CAMPOS DEL ESTUDIANTE ===
        student_frame = ttk.LabelFrame(left_column, text="Campos del Estudiante", padding="10")
        student_frame.grid(row=row_left, column=0, sticky="ew", pady=5)

        ttk.Checkbutton(student_frame, text="Mostrar campo 'Nombre'", variable=self.show_student_name).grid(row=0, column=0, sticky="w", padx=5, pady=2)
        ttk.Checkbutton(student_frame, text="Mostrar campo 'RUT/ID'", variable=self.show_student_id).grid(row=1, column=0, sticky="w", padx=5, pady=2)
        ttk.Checkbutton(student_frame, text="Mostrar campo 'Sección'", variable=self.show_student_section).grid(row=2, column=0, sticky="w", padx=5, pady=2)
        ttk.Checkbutton(student_frame, text="Mostrar campo 'Puntaje'", variable=self.show_exam_score).grid(row=3, column=0, sticky="w", padx=5, pady=2)

        row_left += 1

        # === CONFIGURACIÓN DE FUENTES ===
        font_frame = ttk.LabelFrame(left_column, text="Configuración de Fuentes", padding="10")
        font_frame.grid(row=row_left, column=0, sticky="ew", pady=5)
        font_frame.columnconfigure(1, weight=1)

        ttk.Label(font_frame, text="Fuente:").grid(row=0, column=0, sticky="w", padx=5, pady=2)
        font_combo = ttk.Combobox(font_frame, textvariable=self.font_name, width=20)
        font_combo['values'] = ('Arial', 'Times New Roman', 'Calibri', 'Helvetica', 'Liberation Sans')
        font_combo.grid(row=0, column=1, sticky="w", padx=5, pady=2)

        ttk.Label(font_frame, text="Tamaño título:").grid(row=1, column=0, sticky="w", padx=5, pady=2)
        ttk.Spinbox(font_frame, from_=12, to=24, textvariable=self.title_font_size, width=10).grid(row=1, column=1, sticky="w", padx=5, pady=2)

        ttk.Label(font_frame, text="Tamaño preguntas:").grid(row=2, column=0, sticky="w", padx=5, pady=2)
        ttk.Spinbox(font_frame, from_=8, to=16, textvariable=self.question_font_size, width=10).grid(row=2, column=1, sticky="w", padx=5, pady=2)

        # === COLUMNA DERECHA ===
        row_right = 0

        # === CONFIGURACIÓN DE PÁGINA ===
        page_frame = ttk.LabelFrame(right_column, text="Configuración de Página", padding="10")
        page_frame.grid(row=row_right, column=0, sticky="ew", pady=5)
        page_frame.columnconfigure(1, weight=1)

        ttk.Label(page_frame, text="Tamaño de página:").grid(row=0, column=0, sticky="w", padx=5, pady=2)
        size_combo = ttk.Combobox(page_frame, textvariable=self.page_size, width=20)
        size_combo['values'] = ('A4', 'Carta (Letter)', 'Legal', 'A3')
        size_combo.grid(row=0, column=1, sticky="w", padx=5, pady=2)

        # Márgenes
        ttk.Label(page_frame, text="Márgenes (cm):").grid(row=1, column=0, columnspan=2, sticky="w", padx=5, pady=(10,2))

        margins_frame = ttk.Frame(page_frame)
        margins_frame.grid(row=2, column=0, columnspan=2, sticky="ew", padx=5)

        ttk.Label(margins_frame, text="Superior:").grid(row=0, column=0, sticky="w", padx=2)
        ttk.Spinbox(margins_frame, from_=1.0, to=5.0, increment=0.5, textvariable=self.margin_top, width=8).grid(row=0, column=1, padx=2)

        ttk.Label(margins_frame, text="Inferior:").grid(row=0, column=2, sticky="w", padx=2)
        ttk.Spinbox(margins_frame, from_=1.0, to=5.0, increment=0.5, textvariable=self.margin_bottom, width=8).grid(row=0, column=3, padx=2)

        ttk.Label(margins_frame, text="Izquierdo:").grid(row=1, column=0, sticky="w", padx=2, pady=2)
        ttk.Spinbox(margins_frame, from_=1.0, to=5.0, increment=0.5, textvariable=self.margin_left, width=8).grid(row=1, column=1, padx=2, pady=2)

        ttk.Label(margins_frame, text="Derecho:").grid(row=1, column=2, sticky="w", padx=2, pady=2)
        ttk.Spinbox(margins_frame, from_=1.0, to=5.0, increment=0.5, textvariable=self.margin_right, width=8).grid(row=1, column=3, padx=2, pady=2)

        row_right += 1

        # === LOGO (ESQUEMA BÁSICO) ===
        logo_frame = ttk.LabelFrame(right_column, text="Logo (Próximamente)", padding="10")
        logo_frame.grid(row=row_right, column=0, sticky="ew", pady=5)
        logo_frame.columnconfigure(1, weight=1)

        ttk.Checkbutton(logo_frame, text="Mostrar logo", variable=self.show_logo).grid(row=0, column=0, columnspan=2, sticky="w", padx=5, pady=2)

        ttk.Label(logo_frame, text="Posición:").grid(row=1, column=0, sticky="w", padx=5, pady=2)
        position_combo = ttk.Combobox(logo_frame, textvariable=self.logo_position, width=15)
        position_combo['values'] = ('top_left', 'top_right', 'top_center')
        position_combo.grid(row=1, column=1, sticky="w", padx=5, pady=2)

        # Botón placeholder para cargar logo
        ttk.Button(logo_frame, text="Cargar logo (Próximamente)", state="disabled").grid(row=2, column=0, columnspan=2, pady=5)
        ttk.Label(logo_frame, text="Formatos: PNG, JPG, BMP", foreground="gray").grid(row=3, column=0, columnspan=2, pady=2)

        row_right += 1

        # === HOJA DE RESPUESTAS ===
        answer_frame = ttk.LabelFrame(right_column, text="Configuración Hoja de Respuestas", padding="10")
        answer_frame.grid(row=row_right, column=0, sticky="ew", pady=5)
        answer_frame.columnconfigure(1, weight=1)

        ttk.Label(answer_frame, text="Título hoja respuestas:").grid(row=0, column=0, sticky="w", padx=5, pady=2)
        ttk.Entry(answer_frame, textvariable=self.answer_sheet_title, width=25).grid(row=0, column=1, sticky="ew", padx=5, pady=2)

        ttk.Checkbutton(answer_frame, text="Mostrar estadísticas del examen", variable=self.show_statistics).grid(row=1, column=0, columnspan=2, sticky="w", padx=5, pady=2)
        ttk.Checkbutton(answer_frame, text="Mostrar información detallada", variable=self.show_detailed_info).grid(row=2, column=0, columnspan=2, sticky="w", padx=5, pady=2)
        ttk.Checkbutton(answer_frame, text="Mostrar información de archivos", variable=self.show_file_info).grid(row=3, column=0, columnspan=2, sticky="w", padx=5, pady=2)

        # === VISTA PREVIA (abajo, ancho completo) ===
        preview_frame = ttk.LabelFrame(scrollable_frame, text="Vista Previa", padding="10")
        preview_frame.grid(row=1, column=0, columnspan=2, sticky="ew", padx=5, pady=10)

        self.preview_text = tk.Text(preview_frame, height=6, wrap=tk.WORD, state="disabled")
        preview_scroll = ttk.Scrollbar(preview_frame, orient="vertical", command=self.preview_text.yview)
        self.preview_text.configure(yscrollcommand=preview_scroll.set)

        self.preview_text.grid(row=0, column=0, sticky="ew", padx=5, pady=5)
        preview_scroll.grid(row=0, column=1, sticky="ns", pady=5)

        preview_frame.columnconfigure(0, weight=1)

        ttk.Button(preview_frame, text="Actualizar Vista Previa", command=self.update_page_preview).grid(row=1, column=0, pady=5)

        # Configurar grid del frame principal
        frame.columnconfigure(0, weight=1)
        frame.rowconfigure(1, weight=1)
        left_column.columnconfigure(0, weight=1)
        right_column.columnconfigure(0, weight=1)



    def setup_step5(self):
        frame = self.frames[4]

        # Título
        ttk.Label(frame, text="Paso 4: Generar documentos", font=("Arial", 14, "bold")).grid(row=0, column=0, pady=10, sticky="w")

        # Información de resumen
        ttk.Label(frame, text="Resumen:").grid(row=1, column=0, pady=5, sticky="w")

        # Texto para el resumen
        self.summary_text = tk.Text(frame, width=60, height=10, wrap=tk.WORD)
        self.summary_text.grid(row=2, column=0, padx=5, pady=5, sticky="nsew")

        # Botón para generar documentos
        ttk.Button(frame, text="Generar documentos DOCX", command=self.generate_docx).grid(row=3, column=0, pady=10)

        # Configurar grid
        frame.columnconfigure(0, weight=1)
        frame.rowconfigure(2, weight=1)


    
    def update_page_preview(self):
        """Actualizar vista previa de la configuración de página"""
        self.preview_text.config(state="normal")
        self.preview_text.delete(1.0, tk.END)
        
        preview_content = []
        
        # Información del examen
        if self.institution_name.get():
            preview_content.append(f"INSTITUCIÓN: {self.institution_name.get()}")
        
        preview_content.append(f"TÍTULO: {self.exam_title.get()}")
        
        if self.course_name.get():
            preview_content.append(f"CURSO: {self.course_name.get()}")
        
        if self.exam_date.get():
            preview_content.append(f"FECHA: {self.exam_date.get()}")
        
        preview_content.append("")
        
        # Campos del estudiante
        preview_content.append("CAMPOS DEL ESTUDIANTE:")
        if self.show_student_name.get():
            preview_content.append("  - Nombre: ________________")
        if self.show_student_id.get():
            preview_content.append("  - RUT/ID: ________________")
        if self.show_student_section.get():
            preview_content.append("  - Sección: ________________")
        if self.show_exam_score.get():
            preview_content.append("  - Puntaje: _____ / _____")
        
        preview_content.append("")
        
        # Configuración de fuentes
        preview_content.append(f"FUENTE: {self.font_name.get()}")
        preview_content.append(f"TAMAÑO TÍTULO: {self.title_font_size.get()}pt")
        preview_content.append(f"TAMAÑO PREGUNTAS: {self.question_font_size.get()}pt")
        
        if self.show_logo.get():
            preview_content.append(f"LOGO: Habilitado (posición: {self.logo_position.get()})")
        
        preview_content.append("")
        preview_content.append("HOJA DE RESPUESTAS:")
        preview_content.append(f"  - Título: {self.answer_sheet_title.get()}")
        if self.show_statistics.get():
            preview_content.append("  - Incluir estadísticas")
        if self.show_detailed_info.get():
            preview_content.append("  - Incluir información detallada")
        if self.show_file_info.get():
            preview_content.append("  - Incluir información de archivos")
        
        self.preview_text.insert(tk.END, "\n".join(preview_content))
        self.preview_text.config(state="disabled")

    def get_total_score(self):
        """Calcular puntaje total del examen"""
        if not self.questions:
            return 0
        
        total = sum(question.get('score', 1.0) for question in self.questions)
        return total


    def add_multiple_choice_question(self):
        """Ventana emergente para crear pregunta de alternativas"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Nueva pregunta de alternativas")
        dialog.geometry("600x500")
        dialog.transient(self.root)
        dialog.grab_set()

        # Centrar ventana
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() // 2) - (600 // 2)
        y = (dialog.winfo_screenheight() // 2) - (500 // 2)
        dialog.geometry(f"600x500+{x}+{y}")

        # Variables
        question_text = tk.StringVar()
        num_options = tk.IntVar(value=4)
        correct_option = tk.IntVar(value=0)
        score = tk.StringVar(value="1")
        
        option_vars = []
        option_entries = []

        # Función para actualizar opciones
        def update_options():
            for widget in options_frame.winfo_children():
                widget.destroy()
            
            option_vars.clear()
            option_entries.clear()
            
            for i in range(num_options.get()):
                frame = ttk.Frame(options_frame)
                frame.pack(fill=tk.X, pady=2)
                
                # Radio button para correcta
                ttk.Radiobutton(frame, variable=correct_option, value=i).pack(side=tk.LEFT)
                
                # Letra de opción
                ttk.Label(frame, text=f"{chr(65+i)})", width=3).pack(side=tk.LEFT)
                
                # Entry para texto
                var = tk.StringVar()
                entry = ttk.Entry(frame, textvariable=var, width=50)
                entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
                
                option_vars.append(var)
                option_entries.append(entry)
            
            # Marcar primera como correcta por defecto
            correct_option.set(0)

        # Interfaz
        # Texto de pregunta
        ttk.Label(dialog, text="Texto de la pregunta:").pack(anchor="w", padx=10, pady=(10,2))
        text_entry = ttk.Entry(dialog, textvariable=question_text, width=70)
        text_entry.pack(padx=10, pady=2, fill=tk.X)

        # Configuración
        config_frame = ttk.Frame(dialog)
        config_frame.pack(fill=tk.X, padx=10, pady=5)

        ttk.Label(config_frame, text="Cantidad de opciones:").pack(side=tk.LEFT)
        num_spin = tk.Spinbox(config_frame, from_=2, to=8, textvariable=num_options, width=5, command=update_options)
        num_spin.pack(side=tk.LEFT, padx=5)

        ttk.Label(config_frame, text="Puntaje:").pack(side=tk.LEFT, padx=(20,5))
        ttk.Entry(config_frame, textvariable=score, width=5).pack(side=tk.LEFT)

        # Opciones
        ttk.Label(dialog, text="Opciones (seleccione la correcta):").pack(anchor="w", padx=10, pady=(10,2))
        
        # Frame scrollable para opciones
        canvas = tk.Canvas(dialog, height=200)
        scrollbar = ttk.Scrollbar(dialog, orient="vertical", command=canvas.yview)
        options_frame = ttk.Frame(canvas)
        
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.create_window((0, 0), window=options_frame, anchor="nw")
        options_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

        # Botones
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)

        def save_question():
            if not question_text.get().strip():
                messagebox.showerror("Error", "Debe ingresar el texto de la pregunta")
                return
            
            # Verificar que todas las opciones tengan texto
            options = []
            for i, var in enumerate(option_vars):
                text = var.get().strip()
                if not text:
                    messagebox.showerror("Error", f"La opción {chr(65+i)} no puede estar vacía")
                    return
                options.append({
                    'text': text,
                    'is_correct': i == correct_option.get(),
                    'feedback': ''
                })
            
            # Crear pregunta
            question = {
                'title': 'Pregunta creada',
                'text': question_text.get().strip(),
                'type': 'multiple_choice',
                'options': options,
                'correct_answer': option_vars[correct_option.get()].get().strip(),
                'feedback': {},
                'source_file': 'Creada manualmente',
                'problems': [],
                'score': float(score.get()) if score.get().replace('.', '').isdigit() else 1.0
            }
            
            # Agregar a la lista
            self.questions.append(question)
            self.update_questions_list()
            
            dialog.destroy()
            messagebox.showinfo("Éxito", "Pregunta de alternativas creada correctamente")

        ttk.Button(button_frame, text="Guardar", command=save_question).pack(side=tk.RIGHT, padx=5)
        ttk.Button(button_frame, text="Cancelar", command=dialog.destroy).pack(side=tk.RIGHT, padx=5)

        # Inicializar opciones
        update_options()
        text_entry.focus()

    def add_essay_question(self):
        """Ventana emergente para crear pregunta de desarrollo"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Nueva pregunta de desarrollo")
        dialog.geometry("500x300")
        dialog.transient(self.root)
        dialog.grab_set()

        # Centrar ventana
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() // 2) - (500 // 2)
        y = (dialog.winfo_screenheight() // 2) - (300 // 2)
        dialog.geometry(f"500x300+{x}+{y}")

        # Variables
        question_text = tk.StringVar()
        lines = tk.IntVar(value=5)
        score = tk.StringVar(value="1")

        # Interfaz
        ttk.Label(dialog, text="Texto de la pregunta:").pack(anchor="w", padx=10, pady=(10,2))
        
        # Text widget para pregunta larga
        text_frame = ttk.Frame(dialog)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        text_widget = tk.Text(text_frame, height=8, wrap=tk.WORD)
        text_scrollbar = ttk.Scrollbar(text_frame, orient="vertical", command=text_widget.yview)
        text_widget.configure(yscrollcommand=text_scrollbar.set)
        
        text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        text_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Configuración
        config_frame = ttk.Frame(dialog)
        config_frame.pack(fill=tk.X, padx=10, pady=5)

        ttk.Label(config_frame, text="Líneas para respuesta:").pack(side=tk.LEFT)
        ttk.Spinbox(config_frame, from_=1, to=20, textvariable=lines, width=5).pack(side=tk.LEFT, padx=5)

        ttk.Label(config_frame, text="Puntaje:").pack(side=tk.LEFT, padx=(20,5))
        ttk.Entry(config_frame, textvariable=score, width=5).pack(side=tk.LEFT)

        # Botones
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill=tk.X, padx=10, pady=10)

        def save_question():
            question_content = text_widget.get("1.0", tk.END).strip()
            if not question_content:
                messagebox.showerror("Error", "Debe ingresar el texto de la pregunta")
                return
            
            # Crear pregunta de desarrollo
            question = {
                'title': 'Pregunta de desarrollo',
                'text': question_content,
                'type': 'essay',
                'lines': lines.get(),
                'correct_answer': None,
                'feedback': {},
                'source_file': 'Creada manualmente',
                'problems': [],
                'score': float(score.get()) if score.get().replace('.', '').isdigit() else 1.0
            }
            
            # Agregar a la lista
            self.questions.append(question)
            self.update_questions_list()
            
            dialog.destroy()
            messagebox.showinfo("Éxito", "Pregunta de desarrollo creada correctamente")

        ttk.Button(button_frame, text="Guardar", command=save_question).pack(side=tk.RIGHT, padx=5)
        ttk.Button(button_frame, text="Cancelar", command=dialog.destroy).pack(side=tk.RIGHT, padx=5)

        text_widget.focus()

    def move_question_up(self):
        """Subir pregunta en la lista"""
        selection = self.questions_listbox.curselection()
        if not selection or selection[0] == 0:
            return
        
        index = selection[0]
        # Intercambiar posiciones
        self.questions[index], self.questions[index-1] = self.questions[index-1], self.questions[index]
        
        # Actualizar lista y mantener selección
        self.update_questions_list()
        self.questions_listbox.selection_set(index-1)
        self.show_question_editor(index-1)

    def move_question_down(self):
        """Bajar pregunta en la lista"""
        selection = self.questions_listbox.curselection()
        if not selection or selection[0] >= len(self.questions) - 1:
            return
        
        index = selection[0]
        # Intercambiar posiciones
        self.questions[index], self.questions[index+1] = self.questions[index+1], self.questions[index]
        
        # Actualizar lista y mantener selección
        self.update_questions_list()
        self.questions_listbox.selection_set(index+1)
        self.show_question_editor(index+1)

    def delete_question(self):
        """Eliminar pregunta seleccionada"""
        selection = self.questions_listbox.curselection()
        if not selection:
            return
        
        index = selection[0]
        question = self.questions[index]
        
        # Confirmar eliminación
        result = messagebox.askyesno(
            "Confirmar eliminación",
            f"¿Está seguro de eliminar la pregunta {index+1}?\n\n"
            f"Texto: {question['text'][:50]}..."
        )
        
        if result:
            # Eliminar pregunta
            del self.questions[index]
            
            # Actualizar lista
            self.update_questions_list()
            
            # Limpiar editor
            for widget in self.editor_frame.winfo_children():
                widget.destroy()
            self.question_info_label.config(text="Seleccione una pregunta para editar")
            
            messagebox.showinfo("Éxito", "Pregunta eliminada correctamente")

    def apply_score_to_all(self):
        """Aplicar puntaje a todas las preguntas"""
        try:
            score = float(self.mass_score.get())
        except ValueError:
            messagebox.showerror("Error", "Ingrese un puntaje válido")
            return
        
        for question in self.questions:
            question['score'] = score
        
        # Actualizar editor si hay pregunta seleccionada
        selection = self.questions_listbox.curselection()
        if selection:
            self.show_question_editor(selection[0])
        
        messagebox.showinfo("Éxito", f"Puntaje {score} aplicado a todas las preguntas")

    def apply_score_by_type(self, question_type):
        """Aplicar puntaje a preguntas del tipo especificado"""
        try:
            score = float(self.mass_score.get())
        except ValueError:
            messagebox.showerror("Error", "Ingrese un puntaje válido")
            return
        
        count = 0
        for question in self.questions:
            if question['type'] == question_type:
                question['score'] = score
                count += 1
        
        # Actualizar editor si hay pregunta seleccionada
        selection = self.questions_listbox.curselection()
        if selection:
            self.show_question_editor(selection[0])
        
        type_name = "alternativas" if question_type == "multiple_choice" else "desarrollo"
        messagebox.showinfo("Éxito", f"Puntaje {score} aplicado a {count} pregunta(s) de {type_name}")




    def setup_navigation(self):
        # Crear frame para los botones de navegación
        nav_frame = ttk.Frame(self.root, padding="10")
        nav_frame.grid(row=1, column=0, sticky="ew")

        # Botones de navegación
        self.back_button = ttk.Button(nav_frame, text="Atrás", command=self.go_back, state=tk.DISABLED)
        self.back_button.pack(side=tk.LEFT, padx=5)

        self.next_button = ttk.Button(nav_frame, text="Siguiente", command=self.go_next)
        self.next_button.pack(side=tk.RIGHT, padx=5)

    def show_step(self, step_num):
        # Ocultar todos los frames
        for frame in self.frames:
            frame.grid_remove()

        # Mostrar el frame actual
        self.frames[step_num].grid(row=0, column=0, sticky="nsew")
        self.current_step = step_num

        # Actualizar estado de los botones de navegación
        self.back_button.config(state=tk.NORMAL if step_num > 0 else tk.DISABLED)
        self.next_button.config(state=tk.NORMAL if step_num < len(self.frames) - 1 else tk.DISABLED)

        # Cambiar texto del botón siguiente en el último paso
        if step_num == len(self.frames) - 1:
            self.next_button.config(text="Finalizar", command=self.finish)
        else:
            self.next_button.config(text="Siguiente", command=self.go_next)

    def go_back(self):
        if self.current_step > 0:
            self.show_step(self.current_step - 1)

    def go_next(self):
        if self.current_step < len(self.frames) - 1:
            # Validar antes de avanzar
            if self.current_step == 0 and not self.gift_files:
                messagebox.showerror("Error", "Debe seleccionar al menos un archivo GIFT")
                return

            if self.current_step == 1 and not self.output_dir.get():
                messagebox.showerror("Error", "Debe seleccionar un directorio de salida")
                return

            if self.current_step == 2 and not self.questions:
                messagebox.showerror("Error", "Debe cargar los archivos GIFT primero")
                return

            self.show_step(self.current_step + 1)

            # Si estamos entrando al paso 4, generar el resumen
            if self.current_step == 3:
                self.update_summary()

    def finish(self):
        """Finalizar el wizard"""
        self.root.quit()


    def add_gift_files(self):
        """Abrir selector de archivos para elegir archivos GIFT"""
        try:
            # Configurar para ocultar archivos ocultos en Linux/Unix
            try:
                self.root.tk.call('tk_getOpenFile', '-foobarbaz')
            except tk.TclError:
                pass
            # Configurar variables para ocultar archivos ocultos
            self.root.tk.call('set', '::tk::dialog::file::showHiddenBtn', '1')
            self.root.tk.call('set', '::tk::dialog::file::showHiddenVar', '0')
        except:
            pass

        filenames = filedialog.askopenfilenames(
            title="Seleccionar archivos GIFT",
            filetypes=[
                ("Archivos GIFT y texto", "*.gift *.txt"),
                ("Archivos GIFT", "*.gift"),
                ("Archivos de texto", "*.txt"),
                ("Archivos Markdown", "*.md"),
                ("Todos los archivos", "*.*")
            ]
        )

        for filename in filenames:
            if filename not in self.gift_files:
                self.gift_files.append(filename)

        self.update_files_listbox()

    def remove_selected_file(self):
        """Quitar el archivo seleccionado de la lista"""
        selection = self.files_listbox.curselection()
        if selection:
            index = selection[0]
            del self.gift_files[index]
            self.update_files_listbox()

    def clear_all_files(self):
        """Limpiar toda la lista de archivos"""
        self.gift_files.clear()
        self.update_files_listbox()

    def update_files_listbox(self):
        """Actualizar la lista de archivos en la interfaz"""
        self.files_listbox.delete(0, tk.END)
        for file_path in self.gift_files:
            filename = os.path.basename(file_path)
            self.files_listbox.insert(tk.END, filename)

    def browse_output_dir(self):
        """Abrir selector de directorio para elegir el directorio de salida"""
        directory = filedialog.askdirectory(title="Seleccionar directorio de salida")
        if directory:
            self.output_dir.set(directory)

    def read_file_with_encoding(self, file_path):
        """Leer archivo probando diferentes codificaciones"""
        encodings_to_try = []

        if self.encoding_var.get() == "auto":
            # Intentar múltiples codificaciones en orden de probabilidad
            encodings_to_try = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        else:
            # Usar la codificación seleccionada específicamente
            encodings_to_try = [self.encoding_var.get()]

        last_error = None
        for encoding in encodings_to_try:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()

                # Normalizar algunos caracteres comunes problemáticos
                content = self.normalize_special_characters(content)
                return content

            except UnicodeDecodeError as e:
                last_error = e
                continue
            except Exception as e:
                last_error = e
                continue

        # Si ninguna codificación funcionó, lanzar el último error
        raise Exception(f"No se pudo leer el archivo {os.path.basename(file_path)} con ninguna codificación. Último error: {str(last_error)}")

    def normalize_special_characters(self, text):
        """Normalizar caracteres especiales comunes"""
        # Diccionario de reemplazos para caracteres problemáticos
        replacements = {
            # Superíndices comunes
            '⁰': '⁰', '¹': '¹', '²': '²', '³': '³', '⁴': '⁴', '⁵': '⁵',
            '⁶': '⁶', '⁷': '⁷', '⁸': '⁸', '⁹': '⁹', '⁺': '⁺', '⁻': '⁻',

            # Subíndices comunes
            '₀': '₀', '₁': '₁', '₂': '₂', '₃': '₃', '₄': '₄', '₅': '₅',
            '₆': '₆', '₇': '₇', '₈': '₈', '₉': '₉', '₊': '₊', '₋': '₋',

            # Flechas
            '→': '→', '←': '←', '↔': '↔', '↑': '↑', '↓': '↓',
            '⟶': '→', '⟵': '←', '⟷': '↔',

            # Símbolos matemáticos y científicos
            '±': '±', '×': '×', '÷': '÷', '≤': '≤', '≥': '≥', '≠': '≠',
            '∞': '∞', '∂': '∂', '∫': '∫', '∑': '∑', '√': '√', 'π': 'π',
            'α': 'α', 'β': 'β', 'γ': 'γ', 'δ': 'δ', 'ε': 'ε', 'θ': 'θ',
            'λ': 'λ', 'μ': 'μ', 'ν': 'ν', 'ρ': 'ρ', 'σ': 'σ', 'τ': 'τ',
            'φ': 'φ', 'χ': 'χ', 'ψ': 'ψ', 'ω': 'ω',

            # Grados y unidades
            '°': '°', '℃': '℃', '℉': '℉', 'Ω': 'Ω', 'μ': 'μ',

            # Símbolos químicos comunes
            '⁺': '⁺', '⁻': '⁻', '²⁺': '²⁺', '³⁺': '³⁺',

            # Comillas y apostrofes problemáticos
            '"': '"', '"': '"', ''': "'", ''': "'",
            '–': '-', '—': '-', '…': '...',
        }

        # Aplicar reemplazos
        for old_char, new_char in replacements.items():
            text = text.replace(old_char, new_char)

        return text

    def detect_question_problems(self, question):
        """Detectar problemas en una pregunta"""
        problems = []

        if question['type'] != 'multiple_choice':
            return problems

        # Obtener opciones correctas e incorrectas
        correct_options = [opt for opt in question['options'] if opt['is_correct']]
        incorrect_options = [opt for opt in question['options'] if not opt['is_correct']]

        if not correct_options or not incorrect_options:
            return problems

        correct_text = correct_options[0]['text']

        # 1. Detectar respuesta correcta significativamente más larga
        threshold = self.threshold_var.get()
        avg_incorrect_length = sum(len(opt['text']) for opt in incorrect_options) / len(incorrect_options)

        if len(correct_text) - avg_incorrect_length > threshold:
            problems.append(f"Respuesta correcta {int(len(correct_text) - avg_incorrect_length)} caracteres más larga que promedio")

        # 2. Detectar distractores muy cortos
        short_distractors = [opt for opt in incorrect_options if len(opt['text']) < 10]
        if short_distractors:
            problems.append(f"{len(short_distractors)} distractor(es) muy corto(s) (< 10 caracteres)")

        # 3. Detectar palabras clave problemáticas
        problematic_phrases = [
            "todas las anteriores", "todas las opciones", "todas son correctas",
            "ninguna de las anteriores", "ninguna opción", "ninguna es correcta",
            "solo a", "solo b", "solo c", "solo d",
            "a y b", "a y c", "a y d", "b y c", "b y d", "c y d",
            "a, b y c", "a, b y d", "a, c y d", "b, c y d"
        ]

        for option in question['options']:
            text_lower = option['text'].lower()
            for phrase in problematic_phrases:
                if phrase in text_lower:
                    problems.append(f"Opción contiene frase problemática: '{phrase}'")
                    break

        return problems

    def process_gift_files(self):
        """Procesar todos los archivos GIFT y extraer las preguntas"""
        if not self.gift_files:
            messagebox.showerror("Error", "Debe seleccionar al menos un archivo GIFT primero")
            return

        try:
            all_questions = []
            encoding_issues = []

            # Procesar cada archivo
            for file_path in self.gift_files:
                try:
                    # Usar la nueva función de lectura con manejo de codificaciones
                    content = self.read_file_with_encoding(file_path)

                    # Procesar el contenido GIFT
                    file_questions = self.parse_gift_questions(content, os.path.basename(file_path))
                    all_questions.extend(file_questions)

                except Exception as e:
                    error_msg = f"Error al procesar {os.path.basename(file_path)}: {str(e)}"
                    encoding_issues.append(error_msg)
                    continue

            # Mostrar advertencias de codificación si las hay
            if encoding_issues:
                messagebox.showwarning(
                    "Problemas de codificación",
                    "Se encontraron problemas al leer algunos archivos:\n\n" +
                    "\n".join(encoding_issues) +
                    "\n\nPrueba seleccionar una codificación específica en las opciones."
                )

            # Aplicar orden según configuración
            if self.question_order.get() == "aleatorio":
                random.shuffle(all_questions)

            self.questions = all_questions

            # Detectar problemas en todas las preguntas
            for question in self.questions:
                question['problems'] = self.detect_question_problems(question)

            # Actualizar lista de preguntas
            self.update_questions_list()

            processed_files = len(self.gift_files) - len(encoding_issues)
            messagebox.showinfo("Éxito", f"Se han procesado {len(self.questions)} preguntas de {processed_files} archivo(s)")

        except Exception as e:
            messagebox.showerror("Error", f"Error general al procesar archivos: {str(e)}")


    def update_questions_list(self):
        """Actualizar la lista de preguntas en el panel izquierdo"""
        self.questions_listbox.delete(0, tk.END)

        for i, question in enumerate(self.questions):
            # Determinar estado de la pregunta
            if question.get('problems'):
                status = "[!]"  # En lugar de ⚠️
            else:
                status = "[OK]" # En lugar de ✅

            # Determinar tipo de pregunta
            if question['type'] == 'multiple_choice':
                type_icon = "[A]"
            elif question['type'] == 'essay':
                type_icon = "[D]"
            else:
                type_icon = "[?]"

            # Mostrar puntaje
            score = question.get('score', 1.0)
            score_text = f"({score}pts)"

            # Truncar texto de pregunta para mostrar en lista
            question_text = question['text'][:35] + "..." if len(question['text']) > 35 else question['text']

            # Agregar a la lista
            display_text = f"{status} {type_icon} {i+1}. {question_text} {score_text}"
            self.questions_listbox.insert(tk.END, display_text)

    def on_question_select(self, event):
        """Manejar selección de pregunta en la lista"""
        selection = self.questions_listbox.curselection()
        if not selection:
            return

        question_index = selection[0]
        self.selected_question_index.set(question_index)
        self.show_question_editor(question_index)

    def show_question_editor(self, question_index):
        """Mostrar editor para la pregunta seleccionada"""
        if question_index >= len(self.questions):
            return

        question = self.questions[question_index]

        # Limpiar el editor anterior
        for widget in self.editor_frame.winfo_children():
            widget.destroy()

        # Información de la pregunta
        type_name = "Alternativas" if question['type'] == 'multiple_choice' else "Desarrollo"
        info_text = f"Pregunta {question_index + 1} - {type_name}"
        if question.get('source_file'):
            info_text += f" (Archivo: {question['source_file']})"
        self.question_info_label.config(text=info_text)

        row = 0

        # Mostrar problemas detectados
        if question.get('problems'):
            problems_frame = ttk.LabelFrame(self.editor_frame, text="[!] Problemas detectados", padding="5")
            problems_frame.grid(row=row, column=0, sticky="ew", padx=5, pady=5)
            problems_frame.columnconfigure(0, weight=1)

            for i, problem in enumerate(question['problems']):
                ttk.Label(problems_frame, text=f"* {problem}", foreground="red", wraplength=400).grid(row=i, column=0, sticky="w")

            row += 1

        # Puntaje de la pregunta
        score_frame = ttk.LabelFrame(self.editor_frame, text="Puntaje", padding="5")
        score_frame.grid(row=row, column=0, sticky="ew", padx=5, pady=5)
        
        ttk.Label(score_frame, text="Puntaje:").grid(row=0, column=0, sticky="w")
        score_var = tk.StringVar(value=str(question.get('score', 1.0)))
        score_entry = ttk.Entry(score_frame, textvariable=score_var, width=10)
        score_entry.grid(row=0, column=1, padx=5, sticky="w")
        question['score_var'] = score_var
        
        def update_score():
            try:
                question['score'] = float(score_var.get())
                self.update_questions_list()
                self.questions_listbox.selection_set(question_index)
            except ValueError:
                messagebox.showerror("Error", "Ingrese un puntaje válido")
                score_var.set(str(question.get('score', 1.0)))
        
        ttk.Button(score_frame, text="Actualizar", command=update_score).grid(row=0, column=2, padx=5)
        
        row += 1

        # Editor del texto de la pregunta
        ttk.Label(self.editor_frame, text="Texto de la pregunta:").grid(row=row, column=0, sticky="w", padx=5, pady=(10,2))
        row += 1

        # Para preguntas largas, usar Text widget
        if len(question['text']) > 100 or '\n' in question['text']:
            text_frame = ttk.Frame(self.editor_frame)
            text_frame.grid(row=row, column=0, sticky="ew", padx=5, pady=2)
            text_frame.columnconfigure(0, weight=1)
            
            question_text = tk.Text(text_frame, height=4, wrap=tk.WORD)
            text_scroll = ttk.Scrollbar(text_frame, orient="vertical", command=question_text.yview)
            question_text.configure(yscrollcommand=text_scroll.set)
            
            question_text.grid(row=0, column=0, sticky="ew")
            text_scroll.grid(row=0, column=1, sticky="ns")
            
            question_text.insert("1.0", question['text'])
            question['text_widget'] = question_text
        else:
            question_text_var = tk.StringVar(value=question['text'])
            question_entry = ttk.Entry(self.editor_frame, textvariable=question_text_var, width=60)
            question_entry.grid(row=row, column=0, sticky="ew", padx=5, pady=2)
            question['text_var'] = question_text_var
        
        row += 1

        # Editor específico según el tipo de pregunta
        if question['type'] == 'multiple_choice':
            # Editor de opciones
            ttk.Label(self.editor_frame, text="Opciones de respuesta:").grid(row=row, column=0, sticky="w", padx=5, pady=(10,2))
            row += 1

            option_chars = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']

            for i, option in enumerate(question['options']):
                option_frame = ttk.Frame(self.editor_frame)
                option_frame.grid(row=row, column=0, sticky="ew", padx=5, pady=2)
                option_frame.columnconfigure(1, weight=1)

                # Letra de la opción
                letter_label = ttk.Label(option_frame, text=f"{option_chars[i]})", width=3)
                letter_label.grid(row=0, column=0, padx=(0,5))

                # Campo de texto editable
                option_var = tk.StringVar(value=option['text'])
                option_entry = ttk.Entry(option_frame, textvariable=option_var, width=50)
                option_entry.grid(row=0, column=1, sticky="ew")
                option['text_var'] = option_var

                # Indicador de longitud y correcto
                char_count = len(option['text'])
                correct_text = " [✓ CORRECTA]" if option['is_correct'] else ""
                info_label = ttk.Label(option_frame, text=f"({char_count} caracteres){correct_text}")
                info_label.grid(row=0, column=2, padx=(5,0))
                option['info_label'] = info_label

                # Actualizar contador cuando cambie el texto
                def update_char_count(var, option=option):
                    new_length = len(var.get())
                    correct_text = " [✓ CORRECTA]" if option['is_correct'] else ""
                    option['info_label'].config(text=f"({new_length} caracteres){correct_text}")

                option_var.trace('w', lambda *args, var=option_var: update_char_count(var))

                row += 1

        elif question['type'] == 'essay':
            # Editor para pregunta de desarrollo
            ttk.Label(self.editor_frame, text="Configuración de desarrollo:").grid(row=row, column=0, sticky="w", padx=5, pady=(10,2))
            row += 1
            
            dev_frame = ttk.Frame(self.editor_frame)
            dev_frame.grid(row=row, column=0, sticky="ew", padx=5, pady=2)
            
            ttk.Label(dev_frame, text="Líneas para respuesta:").grid(row=0, column=0, sticky="w")
            lines_var = tk.IntVar(value=question.get('lines', 5))
            lines_spinbox = ttk.Spinbox(dev_frame, from_=1, to=20, textvariable=lines_var, width=5)
            lines_spinbox.grid(row=0, column=1, padx=5, sticky="w")
            question['lines_var'] = lines_var
            
            def update_lines():
                question['lines'] = lines_var.get()
            
            lines_var.trace('w', lambda *args: update_lines())
            
            row += 1

        # Botones de acción para la pregunta
        buttons_frame = ttk.Frame(self.editor_frame)
        buttons_frame.grid(row=row, column=0, sticky="ew", padx=5, pady=10)

        ttk.Button(buttons_frame, text="Guardar cambios", command=lambda: self.save_question_changes(question_index)).pack(side=tk.LEFT, padx=5)
        
        if question['type'] == 'multiple_choice':
            ttk.Button(buttons_frame, text="Reanalizar problemas", command=lambda: self.reanalyze_question(question_index)).pack(side=tk.LEFT, padx=5)

        # Configurar scrolling
        self.editor_frame.columnconfigure(0, weight=1)
        self.editor_canvas.update_idletasks()
        self.editor_canvas.configure(scrollregion=self.editor_canvas.bbox("all"))

    def quick_cleanup(self, text_to_remove):
        """Aplicar limpieza rápida con texto predefinido"""
        self.cleanup_text.set(text_to_remove)
        self.apply_mass_cleanup()

    def apply_mass_cleanup(self):
        """Aplicar limpieza masiva a todas las preguntas"""
        if not self.questions:
            messagebox.showerror("Error", "No hay preguntas cargadas")
            return

        text_to_remove = self.cleanup_text.get().strip()
        if not text_to_remove:
            messagebox.showerror("Error", "Ingrese el texto a eliminar")
            return

        # Confirmar acción
        result = messagebox.askyesno(
            "Confirmar limpieza masiva",
            f"¿Está seguro de eliminar '{text_to_remove}' de todas las preguntas y opciones?\n\n"
            "Esta acción no se puede deshacer."
        )

        if not result:
            return

        changes_made = 0

        # Aplicar limpieza a todas las preguntas
        for question in self.questions:
            # Limpiar texto de pregunta
            if text_to_remove in question['text']:
                question['text'] = question['text'].replace(text_to_remove, '')
                changes_made += 1

            # Limpiar opciones
            if question['type'] == 'multiple_choice':
                for option in question['options']:
                    if text_to_remove in option['text']:
                        option['text'] = option['text'].replace(text_to_remove, '')
                        changes_made += 1

                    # Limpiar retroalimentación si existe
                    if option.get('feedback') and text_to_remove in option['feedback']:
                        option['feedback'] = option['feedback'].replace(text_to_remove, '')
                        changes_made += 1

        # Reanalizar problemas después de la limpieza
        for question in self.questions:
            question['problems'] = self.detect_question_problems(question)

        # Actualizar interfaz
        self.update_questions_list()

        # Si hay una pregunta seleccionada, actualizar el editor
        current_selection = self.questions_listbox.curselection()
        if current_selection:
            self.show_question_editor(current_selection[0])

        # Limpiar campo de entrada
        self.cleanup_text.set("")

        messagebox.showinfo(
            "Limpieza completada",
            f"Limpieza masiva completada.\n\n"
            f"Se realizaron {changes_made} cambios en total.\n"
            f"Los problemas han sido reanalizado automáticamente."
        )


    def reanalyze_question(self, question_index):
        """Reanalizar problemas de la pregunta actual"""
        if question_index >= len(self.questions):
            return

        question = self.questions[question_index]

        # Guardar cambios primero
        self.save_question_changes(question_index)

        # Mostrar editor actualizado
        self.show_question_editor(question_index)
    
    def save_question_changes(self, question_index):
        """Guardar cambios de la pregunta actual"""
        if question_index >= len(self.questions):
            return

        question = self.questions[question_index]

        # Actualizar texto de pregunta
        if hasattr(question, 'text_var'):
            question['text'] = question['text_var'].get()
        elif hasattr(question, 'text_widget'):
            question['text'] = question['text_widget'].get("1.0", tk.END).strip()

        # Actualizar puntaje
        if hasattr(question, 'score_var'):
            try:
                question['score'] = float(question['score_var'].get())
            except ValueError:
                question['score'] = 1.0

        # Actualizar opciones para preguntas de alternativas
        if question['type'] == 'multiple_choice':
            for option in question['options']:
                if hasattr(option, 'text_var'):
                    option['text'] = option['text_var'].get()

        # Actualizar líneas para preguntas de desarrollo
        elif question['type'] == 'essay':
            if hasattr(question, 'lines_var'):
                question['lines'] = question['lines_var'].get()

        # Reanalizar problemas solo para preguntas de alternativas
        if question['type'] == 'multiple_choice':
            question['problems'] = self.detect_question_problems(question)

        # Actualizar lista de preguntas
        self.update_questions_list()

        # Reseleccionar la pregunta actual
        self.questions_listbox.selection_set(question_index)

        messagebox.showinfo("Guardado", "Cambios guardados correctamente")

    def save_as_gift(self):
        """Guardar las preguntas editadas como un nuevo archivo GIFT"""
        if not self.questions:
            messagebox.showerror("Error", "No hay preguntas para guardar")
            return

        # Solicitar nombre del archivo
        filename = filedialog.asksaveasfilename(
            title="Guardar como archivo GIFT",
            defaultextension=".gift",
            filetypes=[("Archivos GIFT", "*.gift"), ("Archivos de texto", "*.txt"), ("Todos los archivos", "*.*")]
        )

        if not filename:
            return

        try:
            # Guardar cambios de la pregunta actual si está siendo editada
            current_selection = self.questions_listbox.curselection()
            if current_selection:
                self.save_question_changes(current_selection[0])

            gift_content = self.questions_to_gift_format()

            with open(filename, 'w', encoding='utf-8') as file:
                file.write(gift_content)

            messagebox.showinfo("Guardado", f"Archivo GIFT guardado correctamente:\n{filename}")

        except Exception as e:
            messagebox.showerror("Error", f"Error al guardar el archivo: {str(e)}")


    def questions_to_gift_format(self):
        """Convertir las preguntas editadas de vuelta al formato GIFT - actualizado para manejar desarrollo"""
        gift_lines = []

        for question in self.questions:
            # Agregar título si existe
            if question.get('title') and question['title'] != "Sin título":
                gift_lines.append(f"::{question['title']}::")

            # Agregar texto de pregunta
            gift_lines.append(question['text'])

            # Agregar opciones si es de opción múltiple
            if question['type'] == 'multiple_choice':
                options_text = "{"

                for option in question['options']:
                    prefix = "=" if option['is_correct'] else "~"
                    option_text = option['text']

                    # Agregar retroalimentación si existe
                    if option.get('feedback'):
                        option_text += f"#{option['feedback']}"

                    options_text += f"{prefix}{option_text}"

                options_text += "}"
                gift_lines.append(options_text)
            
            elif question['type'] == 'essay':
                # Para preguntas de desarrollo, agregar formato especial
                lines = question.get('lines', 5)
                gift_lines.append(f"{{# Pregunta de desarrollo - {lines} líneas}}")

            # Línea en blanco entre preguntas
            gift_lines.append("")

        return "\n".join(gift_lines)

    def apply_randomization(self):
        """Aplicar aleatorización y continuar al siguiente paso"""
        if not self.questions:
            messagebox.showerror("Error", "No hay preguntas para procesar")
            return

        # Guardar cambios de la pregunta actual si está siendo editada
        current_selection = self.questions_listbox.curselection()
        if current_selection:
            self.save_question_changes(current_selection[0])

        # Verificar si hay problemas no resueltos
        questions_with_problems = [q for q in self.questions if q.get('problems')]

        if questions_with_problems:
            result = messagebox.askyesno(
                "Problemas detectados",
                f"Hay {len(questions_with_problems)} pregunta(s) con problemas detectados.\n\n"
                "¿Desea continuar de todos modos?"
            )
            if not result:
                return

        messagebox.showinfo("Procesado", "Preguntas procesadas. Aleatorización aplicada según configuración.")
        self.go_next()

    def parse_gift_questions(self, content, source_file=""):
        """Parsear preguntas en formato GIFT - actualizado para incluir puntajes"""
        # Eliminar comentarios
        content = re.sub(r'//.*?$', '', content, flags=re.MULTILINE)


        # Dividir en preguntas individuales
        question_blocks = re.split(r'\n\s*\n', content)

        questions = []

        for block in question_blocks:
            block = block.strip()
            if not block:
                continue

            try:
                # Extraer el título (si existe)
                title_match = re.match(r'::(.+?)::', block)
                title = title_match.group(1).strip() if title_match else "Sin título"

                # Remover el título del bloque para seguir procesando
                if title_match:
                    block = block[title_match.end():].strip()

                # Extraer el texto de la pregunta (todo hasta la primera respuesta)
                question_text_match = re.match(r'(.*?)\{', block, re.DOTALL)
                if not question_text_match:
                    continue

                question_text = question_text_match.group(1).strip()

                # Extraer las opciones y retroalimentación
                options_block = block[question_text_match.end()-1:]

                # Procesar según el tipo de pregunta
                question_data = {
                    'title': title,
                    'text': question_text,
                    'type': 'unknown',
                    'options': [],
                    'correct_answer': None,
                    'feedback': {},
                    'source_file': source_file,
                    'problems': [],
                    'score': 1.0  # Puntaje por defecto
                }

                # Verifica si es una pregunta de opción múltiple
                if re.match(r'\{', options_block):
                    # Dividir las opciones por ~ o = al inicio, conservando el delimitador
                    options_match = re.split(r'(?=[~=])', options_block.strip()[1:-1].strip())

                    options = []
                    for option_text in options_match:
                        option_text = option_text.strip()
                        if not option_text:
                            continue

                        is_correct = option_text.startswith('=')
                        option_text = option_text[1:].strip()  # Eliminar el prefijo ~ o =

                        # Extraer retroalimentación
                        feedback_text = ""
                        feedback_idx = option_text.find('#')

                        if feedback_idx != -1:
                            feedback_text = option_text[feedback_idx+1:].strip()
                            option_text = option_text[:feedback_idx].strip()

                        options.append({
                            'text': option_text,
                            'is_correct': is_correct,
                            'feedback': feedback_text
                        })

                        if is_correct:
                            question_data['correct_answer'] = option_text

                    question_data['type'] = 'multiple_choice'
                    question_data['options'] = options

                # Si hay al menos texto de pregunta, la añadimos
                if question_data['text']:
                    questions.append(question_data)

            except Exception as e:
                print(f"Error al procesar una pregunta en {source_file}: {str(e)}")
                continue

        return questions

    def update_summary(self):
        """Actualizar el resumen en el paso 4"""
        self.summary_text.delete(1.0, tk.END)

        if not self.questions:
            self.summary_text.insert(tk.END, "No hay preguntas para procesar.")
            return

        # Información de archivos
        self.summary_text.insert(tk.END, f"Archivos procesados: {len(self.gift_files)}\n")
        for i, file_path in enumerate(self.gift_files, 1):
            filename = os.path.basename(file_path)
            file_questions = sum(1 for q in self.questions if q['source_file'] == filename)
            self.summary_text.insert(tk.END, f"  {i}. {filename} ({file_questions} preguntas)\n")

        self.summary_text.insert(tk.END, f"\nDirectorio de salida: {self.output_dir.get()}\n\n")

        # Información de preguntas
        self.summary_text.insert(tk.END, f"Total de preguntas: {len(self.questions)}\n")
        self.summary_text.insert(tk.END, f"Preguntas de opción múltiple: {sum(1 for q in self.questions if q['type'] == 'multiple_choice')}\n")
        self.summary_text.insert(tk.END, f"Orden de preguntas: {'Aleatorio' if self.question_order.get() == 'aleatorio' else 'Por archivos'}\n")
        self.summary_text.insert(tk.END, f"Opciones aleatorizadas: {'Sí' if self.randomize_options.get() else 'No'}\n")

        # Información de problemas detectados
        questions_with_problems = [q for q in self.questions if q.get('problems')]
        self.summary_text.insert(tk.END, f"Preguntas con problemas detectados: {len(questions_with_problems)}\n\n")

        # Analizar el número de opciones por pregunta
        option_counts = [len(q['options']) for q in self.questions if q['type'] == 'multiple_choice']
        if option_counts:
            counter = collections.Counter(option_counts)
            self.summary_text.insert(tk.END, "Distribución de opciones por pregunta:\n")
            for count, frequency in sorted(counter.items()):
                self.summary_text.insert(tk.END, f"  - Preguntas con {count} opciones: {frequency}\n")

        self.summary_text.insert(tk.END, "\nSe generarán dos archivos DOCX:\n")
        self.summary_text.insert(tk.END, "1. Examen.docx - Contiene las preguntas y hoja de respuestas\n")
        self.summary_text.insert(tk.END, "2. Respuestas.docx - Contiene las respuestas correctas con estadísticas\n")

    def generate_docx(self):
        """Generar los documentos DOCX"""
        if not self.questions:
            messagebox.showerror("Error", "No hay preguntas para generar los documentos")
            return

        if not self.output_dir.get() or not os.path.isdir(self.output_dir.get()):
            messagebox.showerror("Error", "El directorio de salida no es válido")
            return

        try:
            # IMPORTANTE: Generar PRIMERO el documento de examen para llenar self.answer_keys
            exam_doc = self.create_exam_document()

            # LUEGO generar el documento de respuestas (que ya tendrá las estadísticas)
            answers_doc = self.create_answers_document()

            # Generar nombres de archivo con timestamp si ya existen
            import datetime
            
            exam_filename = "Examen.docx"
            answers_filename = "Respuestas.docx"
            
            exam_path = os.path.join(self.output_dir.get(), exam_filename)
            answers_path = os.path.join(self.output_dir.get(), answers_filename)
            
            # Si los archivos ya existen, agregar timestamp
            if os.path.exists(exam_path) or os.path.exists(answers_path):
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                exam_filename = f"Examen_{timestamp}.docx"
                answers_filename = f"Respuestas_{timestamp}.docx"
                exam_path = os.path.join(self.output_dir.get(), exam_filename)
                answers_path = os.path.join(self.output_dir.get(), answers_filename)

            # Guardar documentos
            exam_doc.save(exam_path)
            answers_doc.save(answers_path)

            messagebox.showinfo(
                "Éxito",
                f"Documentos generados correctamente:\n\n"
                f"- Examen: {exam_path}\n"
                f"- Respuestas: {answers_path}"
            )

        except Exception as e:
            messagebox.showerror("Error", f"Error al generar los documentos: {str(e)}")

    def create_exam_document(self):
        """Crear documento de examen con configuraciones del Paso 4"""
        doc = Document()

        # Configurar página según configuraciones
        section = doc.sections[0]
        
        # Configurar tamaño de página
        if self.page_size.get() == "A4":
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
        elif self.page_size.get() == "Carta (Letter)":
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        elif self.page_size.get() == "Legal":
            section.page_width = Inches(8.5)
            section.page_height = Inches(14)
        elif self.page_size.get() == "A3":
            section.page_width = Inches(11.69)
            section.page_height = Inches(16.54)
        
        # Configurar márgenes (convertir cm a inches: cm * 0.393701)
        section.top_margin = Inches(self.margin_top.get() * 0.393701)
        section.bottom_margin = Inches(self.margin_bottom.get() * 0.393701)
        section.left_margin = Inches(self.margin_left.get() * 0.393701)
        section.right_margin = Inches(self.margin_right.get() * 0.393701)

        # Configurar estilos con fuente seleccionada
        style = doc.styles['Normal']
        style.font.name = self.font_name.get()
        style.font.size = Pt(self.question_font_size.get())

        # Encabezado institucional
        if self.institution_name.get():
            institution_p = doc.add_paragraph()
            institution_run = institution_p.add_run(self.institution_name.get())
            institution_run.font.name = self.font_name.get()
            institution_run.font.size = Pt(self.title_font_size.get() - 2)
            institution_run.bold = True
            institution_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Título del examen
        title = doc.add_paragraph()
        title_run = title.add_run(self.exam_title.get())
        title_run.font.name = self.font_name.get()
        title_run.font.size = Pt(self.title_font_size.get())
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Información del curso y fecha
        if self.course_name.get() or self.exam_date.get():
            info_p = doc.add_paragraph()
            info_text = []
            if self.course_name.get():
                info_text.append(f"Curso: {self.course_name.get()}")
            if self.exam_date.get():
                info_text.append(f"Fecha: {self.exam_date.get()}")
            
            info_run = info_p.add_run(" | ".join(info_text))
            info_run.font.name = self.font_name.get()
            info_run.font.size = Pt(self.question_font_size.get())
            info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Información del estudiante (solo campos habilitados)
        student_fields = []
        if self.show_student_name.get():
            student_fields.append("Nombre: ________________________________")
        if self.show_student_id.get():
            student_fields.append("RUT/ID: ________________________________")
        if self.show_student_section.get():
            student_fields.append("Sección: ________________________________")
        if self.show_exam_score.get():
            total_score = self.get_total_score()
            student_fields.append(f"Puntaje: ______ / {total_score}")

        if student_fields:
            for field in student_fields:
                field_p = doc.add_paragraph(field)
                field_p.style.font.name = self.font_name.get()
                field_p.style.font.size = Pt(self.question_font_size.get())

        doc.add_paragraph()

        # Limpiar el diccionario de respuestas
        self.answer_keys = {}

        # Analizar el número de opciones para cada pregunta
        options_per_question = {}
        for i, question in enumerate(self.questions, 1):
            if question['type'] == 'multiple_choice':
                options_per_question[i] = len(question['options'])
            elif question['type'] == 'essay':
                options_per_question[i] = 0  # Sin opciones para desarrollo

        # Encontrar el máximo número de opciones por pregunta
        max_options = max([opt for opt in options_per_question.values() if opt > 0], default=4)

        # Calcular la distribución óptima de letras correctas
        option_chars = ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h']
        distributions = {}

        # Agrupar preguntas por número de opciones (solo alternativas)
        for q_num, num_options in options_per_question.items():
            if num_options > 0:  # Solo preguntas de alternativas
                if num_options not in distributions:
                    distributions[num_options] = []
                distributions[num_options].append(q_num)

        # Para cada grupo, asignar letras correctas de manera equitativa
        for num_options, q_numbers in distributions.items():
            valid_chars = option_chars[:num_options]
            target_answers = []
            for char in valid_chars:
                target_answers.extend([char] * (len(q_numbers) // len(valid_chars)))

            remaining = len(q_numbers) - len(target_answers)
            if remaining > 0:
                target_answers.extend(random.sample(valid_chars, remaining))

            random.shuffle(target_answers)

            for i, q_num in enumerate(q_numbers):
                self.answer_keys[q_num] = target_answers[i]

        # Preguntas
        for i, question in enumerate(self.questions, 1):
            # Añadir texto de la pregunta
            question_p = doc.add_paragraph(f"{i}. {question['text']}")
            question_p.style.font.name = self.font_name.get()
            question_p.style.font.size = Pt(self.question_font_size.get())

            # Si es de opción múltiple, añadir opciones
            if question['type'] == 'multiple_choice':
                options = question['options'].copy()
                num_options = len(options)

                # Reorganizar opciones según el algoritmo mejorado
                if self.randomize_options.get() and i in self.answer_keys:
                    correct_option = next((opt for opt in options if opt['is_correct']), None)
                    if correct_option:
                        target_position = ord(self.answer_keys[i]) - ord('a')
                        options.remove(correct_option)
                        random.shuffle(options)
                        if target_position < len(options) + 1:
                            options.insert(target_position, correct_option)
                        else:
                            options.append(correct_option)
                else:
                    random.shuffle(options)
                    for j, opt in enumerate(options):
                        if opt['is_correct'] and i not in self.answer_keys:
                            self.answer_keys[i] = option_chars[j] if j < len(option_chars) else 'X'

                # Añadir opciones
                for j, opt in enumerate(options):
                    if j < len(option_chars):
                        option_p = doc.add_paragraph(f"   {option_chars[j].upper()}) {opt['text']}")
                        option_p.style.font.name = self.font_name.get()
                        option_p.style.font.size = Pt(self.question_font_size.get())

            elif question['type'] == 'essay':
                # Añadir líneas para respuesta de desarrollo
                lines = question.get('lines', 5)
                for _ in range(lines):
                    line_p = doc.add_paragraph("_" * 80)
                    line_p.style.font.name = self.font_name.get()
                    line_p.style.font.size = Pt(self.question_font_size.get())

            # Espacio entre preguntas
            doc.add_paragraph()

        # Agregar salto de página antes de la hoja de respuestas
        doc.add_page_break()

        # Título de la hoja de respuestas
        answer_title = doc.add_paragraph()
        answer_title_run = answer_title.add_run(self.answer_sheet_title.get())
        answer_title_run.font.name = self.font_name.get()
        answer_title_run.font.size = Pt(self.title_font_size.get())
        answer_title_run.bold = True
        answer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Información del estudiante en hoja de respuestas
        if student_fields:
            doc.add_paragraph("Marque con una X la respuesta correcta para cada pregunta.")
            doc.add_paragraph()
            
            for field in student_fields:
                field_p = doc.add_paragraph(field)
                field_p.style.font.name = self.font_name.get()
                field_p.style.font.size = Pt(self.question_font_size.get())

        doc.add_paragraph()

        # Crear tabla para la hoja de respuestas
        num_questions = len(self.questions)
        rows_per_column = (num_questions + 1) // 2

        cols_per_group = max_options + 1
        total_cols = cols_per_group * 2

        table = doc.add_table(rows=rows_per_column + 1, cols=total_cols)
        table.style = 'Table Grid'

        # Encabezados
        header_letters = [chr(65 + i) for i in range(max_options)]

        cell = table.cell(0, 0)
        cell.text = "Preg"

        for i, letter in enumerate(header_letters):
            cell = table.cell(0, i + 1)
            cell.text = letter

        if num_questions > rows_per_column:
            cell = table.cell(0, cols_per_group)
            cell.text = "Preg"

            for i, letter in enumerate(header_letters):
                cell = table.cell(0, cols_per_group + i + 1)
                cell.text = letter

        # Llenar la tabla
        shading_black_xml = parse_xml(f'<w:shd {nsdecls("w")} w:fill="000000"/>')

        for i in range(num_questions):
            q_num = i + 1
            row_idx = (i % rows_per_column) + 1
            question = self.questions[i]

            if q_num in options_per_question:
                q_options = options_per_question[q_num]
            else:
                q_options = 0

            if i < rows_per_column:  # Primera columna
                cell = table.cell(row_idx, 0)
                cell.text = str(q_num)

                # Manejar celdas según tipo de pregunta
                if question['type'] == 'essay':
                    # Pregunta de desarrollo: todas las celdas negras
                    for j in range(1, max_options + 1):
                        cell = table.cell(row_idx, j)
                        cell._element.get_or_add_tcPr().append(shading_black_xml)
                        cell.text = ""
                else:
                    # Pregunta de alternativas: sombrear celdas que superan el número de opciones
                    for j in range(q_options + 1, max_options + 1):
                        cell = table.cell(row_idx, j)
                        cell._element.get_or_add_tcPr().append(shading_black_xml)
                        cell.text = ""
            else:  # Segunda columna
                cell = table.cell(row_idx, cols_per_group)
                cell.text = str(q_num)

                if question['type'] == 'essay':
                    # Pregunta de desarrollo: todas las celdas negras
                    for j in range(1, max_options + 1):
                        cell = table.cell(row_idx, cols_per_group + j)
                        cell._element.get_or_add_tcPr().append(shading_black_xml)
                        cell.text = ""
                else:
                    # Pregunta de alternativas: sombrear celdas que superan el número de opciones
                    for j in range(q_options + 1, max_options + 1):
                        cell = table.cell(row_idx, cols_per_group + j)
                        cell._element.get_or_add_tcPr().append(shading_black_xml)
                        cell.text = ""

        # Ajustar ancho de columnas
        for cell in table.columns[0].cells:
            cell.width = Inches(0.5)

        if cols_per_group < len(table.columns):
            for cell in table.columns[cols_per_group].cells:
                cell.width = Inches(0.5)

        return doc


    def create_answers_document(self):
        """Crear documento de respuestas con configuraciones del Paso 4"""
        doc = Document()

        # Configurar página según configuraciones (igual que examen)
        section = doc.sections[0]
        
        if self.page_size.get() == "A4":
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
        elif self.page_size.get() == "Carta (Letter)":
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        elif self.page_size.get() == "Legal":
            section.page_width = Inches(8.5)
            section.page_height = Inches(14)
        elif self.page_size.get() == "A3":
            section.page_width = Inches(11.69)
            section.page_height = Inches(16.54)
        
        section.top_margin = Inches(self.margin_top.get() * 0.393701)
        section.bottom_margin = Inches(self.margin_bottom.get() * 0.393701)
        section.left_margin = Inches(self.margin_left.get() * 0.393701)
        section.right_margin = Inches(self.margin_right.get() * 0.393701)

        # Configurar estilos
        style = doc.styles['Normal']
        style.font.name = self.font_name.get()
        style.font.size = Pt(self.question_font_size.get())

        # Encabezado
        if self.institution_name.get():
            institution_p = doc.add_paragraph()
            institution_run = institution_p.add_run(self.institution_name.get())
            institution_run.font.name = self.font_name.get()
            institution_run.font.size = Pt(self.title_font_size.get() - 2)
            institution_run.bold = True
            institution_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Título
        title = doc.add_paragraph()
        title_run = title.add_run("RESPUESTAS CORRECTAS")
        title_run.font.name = self.font_name.get()
        title_run.font.size = Pt(self.title_font_size.get())
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if self.course_name.get() or self.exam_date.get():
            info_p = doc.add_paragraph()
            info_text = []
            if self.course_name.get():
                info_text.append(f"Curso: {self.course_name.get()}")
            if self.exam_date.get():
                info_text.append(f"Fecha: {self.exam_date.get()}")
            
            info_run = info_p.add_run(" | ".join(info_text))
            info_run.font.name = self.font_name.get()
            info_run.font.size = Pt(self.question_font_size.get())
            info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Mostrar estadísticas según configuración
        if self.show_statistics.get():
            stats_p = doc.add_heading('Estadísticas del examen', level=2)
            stats_p.style.font.name = self.font_name.get()

            total_score = self.get_total_score()
            multiple_choice_count = sum(1 for q in self.questions if q['type'] == 'multiple_choice')
            essay_count = sum(1 for q in self.questions if q['type'] == 'essay')

            doc.add_paragraph(f"Total de preguntas: {len(self.questions)}")
            doc.add_paragraph(f"Preguntas de alternativas: {multiple_choice_count}")
            doc.add_paragraph(f"Preguntas de desarrollo: {essay_count}")
            doc.add_paragraph(f"Puntaje total: {total_score} puntos")

        if self.show_file_info.get() and len(self.gift_files) > 1:
            files_p = doc.add_heading('Archivos procesados', level=2)
            files_p.style.font.name = self.font_name.get()
            
            for i, file_path in enumerate(self.gift_files, 1):
                filename = os.path.basename(file_path)
                file_questions = sum(1 for q in self.questions if q['source_file'] == filename)
                doc.add_paragraph(f"{i}. {filename} ({file_questions} preguntas)")

        if self.show_detailed_info.get():
            # Análisis de distribución de respuestas correctas
            if self.answer_keys:
                answer_counter = collections.Counter(self.answer_keys.values())
                dist_p = doc.add_heading('Distribución de alternativas correctas', level=2)
                dist_p.style.font.name = self.font_name.get()

                for letter, count in sorted(answer_counter.items()):
                    doc.add_paragraph(f"Alternativa {letter.upper()}: {count} preguntas")

        doc.add_paragraph()

        # TABLA DE RESPUESTAS CORRECTAS (IDÉNTICA AL EXAMEN PERO CON RESPUESTAS)
        table_title = doc.add_heading('Hoja de Respuestas Correctas', level=2)
        table_title.style.font.name = self.font_name.get()
        
        # Recrear la misma tabla que en el examen
        num_questions = len(self.questions)
        rows_per_column = (num_questions + 1) // 2
        
        # Encontrar máximo de opciones
        max_options = 4
        for question in self.questions:
            if question['type'] == 'multiple_choice':
                max_options = max(max_options, len(question['options']))

        cols_per_group = max_options + 1
        total_cols = cols_per_group * 2

        table = doc.add_table(rows=rows_per_column + 1, cols=total_cols)
        table.style = 'Table Grid'

        # Encabezados idénticos
        header_letters = [chr(65 + i) for i in range(max_options)]

        cell = table.cell(0, 0)
        cell.text = "Preg"

        for i, letter in enumerate(header_letters):
            cell = table.cell(0, i + 1)
            cell.text = letter

        if num_questions > rows_per_column:
            cell = table.cell(0, cols_per_group)
            cell.text = "Preg"

            for i, letter in enumerate(header_letters):
                cell = table.cell(0, cols_per_group + i + 1)
                cell.text = letter

        # Llenar tabla con respuestas correctas
        shading_black_xml = parse_xml(f'<w:shd {nsdecls("w")} w:fill="000000"/>')

        for i in range(num_questions):
            q_num = i + 1
            row_idx = (i % rows_per_column) + 1
            question = self.questions[i]

            if i < rows_per_column:  # Primera columna
                cell = table.cell(row_idx, 0)
                cell.text = str(q_num)

                if question['type'] == 'essay':
                    # Pregunta de desarrollo: todas las celdas negras
                    for j in range(1, max_options + 1):
                        cell = table.cell(row_idx, j)
                        cell._element.get_or_add_tcPr().append(shading_black_xml)
                        cell.text = ""
                else:
                    # Pregunta de alternativas: mostrar respuesta correcta
                    num_options = len(question['options'])
                    correct_letter = self.answer_keys.get(q_num, 'X')
                    correct_index = ord(correct_letter.upper()) - ord('A')

                    for j in range(1, max_options + 1):
                        cell = table.cell(row_idx, j)
                        option_index = j - 1
                        
                        if option_index >= num_options:
                            # Opción no existe: celda negra
                            cell._element.get_or_add_tcPr().append(shading_black_xml)
                            cell.text = ""
                        elif option_index == correct_index:
                            # Respuesta correcta: mostrar letra
                            cell.text = correct_letter.upper()
                        else:
                            # Respuesta incorrecta: celda negra
                            cell._element.get_or_add_tcPr().append(shading_black_xml)
                            cell.text = ""
            else:  # Segunda columna
                cell = table.cell(row_idx, cols_per_group)
                cell.text = str(q_num)

                if question['type'] == 'essay':
                    # Pregunta de desarrollo: todas las celdas negras
                    for j in range(1, max_options + 1):
                        cell = table.cell(row_idx, cols_per_group + j)
                        cell._element.get_or_add_tcPr().append(shading_black_xml)
                        cell.text = ""
                else:
                    # Pregunta de alternativas: mostrar respuesta correcta
                    num_options = len(question['options'])
                    correct_letter = self.answer_keys.get(q_num, 'X')
                    correct_index = ord(correct_letter.upper()) - ord('A')

                    for j in range(1, max_options + 1):
                        cell = table.cell(row_idx, cols_per_group + j)
                        option_index = j - 1
                        
                        if option_index >= num_options:
                            # Opción no existe: celda negra
                            cell._element.get_or_add_tcPr().append(shading_black_xml)
                            cell.text = ""
                        elif option_index == correct_index:
                            # Respuesta correcta: mostrar letra
                            cell.text = correct_letter.upper()
                        else:
                            # Respuesta incorrecta: celda negra
                            cell._element.get_or_add_tcPr().append(shading_black_xml)
                            cell.text = ""

        doc.add_paragraph()

        # Detalles de preguntas (según configuración)
        if self.show_detailed_info.get():
            details_p = doc.add_heading('Detalles de las preguntas', level=2)
            details_p.style.font.name = self.font_name.get()

            for i, question in enumerate(self.questions, 1):
                if len(self.gift_files) > 1:
                    source_p = doc.add_paragraph()
                    source_run = source_p.add_run(f"[Archivo: {question['source_file']}]")
                    source_run.italic = True

                question_p = doc.add_paragraph()
                question_run = question_p.add_run(f"{i}. {question['text']}")
                question_run.bold = True

                if question['type'] == 'multiple_choice':
                    correct_letter = self.answer_keys.get(i, "?")
                    correct_options = [opt for opt in question['options'] if opt['is_correct']]

                    if correct_options:
                        correct_answer = correct_options[0]['text']
                        score = question.get('score', 1.0)
                        doc.add_paragraph(f"   Respuesta correcta: {correct_letter.upper()}) {correct_answer}")
                        doc.add_paragraph(f"   Puntaje: {score} punto(s)")

                        if correct_options[0]['feedback']:
                            doc.add_paragraph(f"   Retroalimentación: {correct_options[0]['feedback']}")
                    else:
                        doc.add_paragraph("   Respuesta correcta: No encontrada")
                
                elif question['type'] == 'essay':
                    score = question.get('score', 1.0)
                    lines = question.get('lines', 5)
                    doc.add_paragraph(f"   Tipo: Pregunta de desarrollo")
                    doc.add_paragraph(f"   Líneas asignadas: {lines}")
                    doc.add_paragraph(f"   Puntaje: {score} punto(s)")

                doc.add_paragraph()

        return doc

def main():
    root = tk.Tk()
    app = GiftToDocxConverter(root)
    root.mainloop()

if __name__ == "__main__":
    main()
